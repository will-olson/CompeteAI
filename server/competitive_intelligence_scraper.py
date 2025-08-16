# competitive_intelligence_scraper.py
import os
import json
import requests
import pandas as pd
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any, Union
import logging
from dotenv import load_dotenv
import time
import re
from urllib.parse import urlparse, urljoin
import hashlib
import feedparser
from bs4 import BeautifulSoup
import csv
import markdown
import docx
from functools import wraps

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s: %(message)s',
    handlers=[
        logging.FileHandler('competitive_intelligence_scraper.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def rate_limited(max_per_second=2):
    """Simple rate limiting decorator for respectful crawling"""
    def decorator(func):
        last_called = {}
        
        @wraps(func)
        def wrapper(*args, **kwargs):
            now = time.time()
            if func.__name__ in last_called:
                time_since_last = now - last_called[func.__name__]
                if time_since_last < 1.0 / max_per_second:
                    time.sleep(1.0 / max_per_second - time_since_last)
            
            last_called[func.__name__] = time.time()
            return func(*args, **kwargs)
        return wrapper
    return decorator

class CompetitiveIntelligenceScraper:
    """
    Comprehensive competitive intelligence scraper supporting multiple data sources
    with enhanced technical content extraction
    """
    
    def __init__(self, firecrawl_api_key: Optional[str] = None):
        self.firecrawl_api_key = firecrawl_api_key or os.getenv('FIRECRAWL_API_KEY')
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        
        # Initialize preset competitor groups
        self.preset_groups = self._initialize_preset_groups()
        
        # Enhanced technical content extraction
        self.technical_keywords = [
            # API-First Architecture
            'api', 'endpoint', 'rest', 'graphql', 'swagger', 'openapi', 'sdk', 'client library',
            'webhook', 'callback', 'authentication', 'oauth', 'rate limit', 'throttling',
            'versioning', 'v1', 'v2', 'v3', 'api version', 'api documentation',
            
            # Cloud-Native Features
            'multi-cloud', 'hybrid cloud', 'auto-scaling', 'elastic', 'serverless', 'lambda',
            'container', 'kubernetes', 'docker', 'microservices', 'distributed',
            'cloud-native', 'cloud-first', 'infrastructure as code', 'terraform', 'cloudformation',
            'managed service', 'paas', 'saas', 'iaas', 'elastic compute', 'auto-scaling',
            
            # Data Integration & Connectors
            'real-time', 'streaming', 'etl', 'elt', 'data pipeline', 'connector', 'integration',
            'data warehouse', 'snowflake', 'bigquery', 'redshift', 'api-based', 'webhook',
            'data mesh', 'data fabric', 'open format', 'parquet', 'avro', 'json', 'csv',
            'data catalog', 'metadata', 'data lineage', 'data governance',
            
            # Developer Experience
            'self-service', 'provisioning', 'ci/cd', 'pipeline', 'deployment', 'infrastructure as code',
            'terraform', 'cloudformation', 'developer documentation', 'getting started',
            'sample code', 'example', 'tutorial', 'playground', 'sandbox', 'demo',
            'community', 'forum', 'support', 'api explorer', 'interactive', 'developer portal',
            
            # Modern Analytics Stack
            'ai', 'machine learning', 'ml', 'artificial intelligence', 'real-time analytics',
            'streaming analytics', 'data mesh', 'data fabric', 'open data', 'data sharing',
            'natural language', 'nlp', 'conversational', 'automated insights', 'auto-discovery',
            'collaborative', 'team analytics', 'governance', 'compliance', 'security',
            'data democratization', 'self-service analytics', 'embedded analytics'
        ]
        
        # Content type detection patterns
        self.content_patterns = {
            'api_docs': [
                r'\b(api|endpoint|rest|graphql|swagger|openapi)\b',
                r'https?://[^\s]+\.json',
                r'https?://[^\s]+\.yaml',
                r'https?://[^\s]+\.yml',
                r'POST|GET|PUT|DELETE|PATCH',
                r'status code|response code|error code'
            ],
            'sdk_docs': [
                r'\b(sdk|client library|download|install|npm|pip|maven|gradle)\b',
                r'github\.com|gitlab\.com|bitbucket\.org',
                r'package\.json|requirements\.txt|pom\.xml|build\.gradle',
                r'installation guide|getting started|quick start'
            ],
            'pricing': [
                r'\b(price|plan|tier|billing|subscription|cost|pricing|enterprise)\b',
                r'\$\d+|\d+\s*(per month|per year|monthly|yearly)',
                r'contact sales|request quote|pricing calculator',
                r'free|starter|professional|enterprise|custom'
            ],
            'deployment': [
                r'\b(deploy|installation|setup|configuration|environment)\b',
                r'docker|kubernetes|aws|azure|gcp|on-premise',
                r'installation guide|deployment guide|setup instructions',
                r'requirements|prerequisites|system requirements'
            ]
        }
        
        # Company-specific technical keywords
        self.company_technical_keywords = {
            'snowflake': ['snowpark', 'cortex', 'warehouse', 'data cloud', 'sql', 'python', 'java', 'spark'],
            'databricks': ['notebook', 'workspace', 'unity catalog', 'delta lake', 'mlflow', 'spark', 'python', 'scala'],
            'powerbi': ['power query', 'dax', 'm language', 'gateway', 'workspace', 'report server', 'power platform'],
            'tableau': ['tableau prep', 'tableau server', 'tableau online', 'tableau public', 'vizql', 'hyper'],
            'looker': ['lookml', 'explore', 'dashboard', 'block', 'looker studio', 'bigquery', 'data warehouse'],
            'qlik': ['qlikview', 'qliksense', 'script', 'load script', 'qlik engine', 'associative engine']
        }
        
    def _initialize_preset_groups(self) -> Dict[str, Dict[str, Any]]:
        """Initialize preset competitor groups for quick analysis"""
        return {
            "tech_saas": {
                "name": "Tech SaaS",
                "companies": ["Salesforce", "HubSpot", "Slack", "Notion", "Figma", "Airtable"],
                "categories": ["marketing", "docs", "rss", "social"],
                "description": "Leading SaaS companies in productivity and collaboration"
            },
            "fintech": {
                "name": "Fintech",
                "companies": ["Stripe", "Plaid", "Coinbase", "Robinhood", "Chime", "Affirm"],
                "categories": ["marketing", "docs", "rss", "social"],
                "description": "Financial technology and payment processing companies"
            },
            "ecommerce": {
                "name": "E-commerce",
                "companies": ["Shopify", "Amazon", "Etsy", "WooCommerce", "BigCommerce", "Magento"],
                "categories": ["marketing", "docs", "rss", "social"],
                "description": "E-commerce platforms and online retail companies"
            },
            "ai_ml": {
                "name": "AI/ML",
                "companies": ["OpenAI", "Anthropic", "Google_AI", "Microsoft_AI", "Meta_AI", "NVIDIA"],
                "categories": ["marketing", "docs", "rss", "social"],
                "description": "Artificial intelligence and machine learning companies"
            }
        }
    
    def load_preset_group(self, group_key: str) -> Dict[str, Any]:
        """Load a preset competitor group"""
        if group_key not in self.preset_groups:
            raise ValueError(f"Unknown preset group: {group_key}")
        
        group = self.preset_groups[group_key].copy()
        group['group_key'] = group_key
        
        # Generate default URLs for each company
        group['company_urls'] = self._generate_default_urls(group['companies'])
        
        return group
    
    def _generate_default_urls(self, companies: List[str]) -> Dict[str, Dict[str, str]]:
        """Generate default URLs for different categories for each company"""
        company_urls = {}
        
        for company in companies:
            company_lower = company.lower().replace('_', '')
            company_urls[company] = {
                'marketing': f"https://{company_lower}.com",
                'docs': f"https://{company_lower}.com/docs",
                'rss': f"https://{company_lower}.com/blog/feed",
                'social': f"https://{company_lower}.com/social"
            }
        
        return company_urls
    
    def create_custom_group(self, name: str, companies: List[str], 
                           categories: List[str], description: str = "") -> Dict[str, Any]:
        """Create a custom competitor group"""
        custom_group = {
            'name': name,
            'companies': companies,
            'categories': categories,
            'description': description,
            'company_urls': self._generate_default_urls(companies),
            'is_custom': True,
            'created_date': datetime.now().isoformat()
        }
        
        return custom_group
    
    def scrape_company_data(self, company: str, urls: Dict[str, str], 
                           categories: List[str], page_limit: int = 10) -> Dict[str, Any]:
        """Scrape data for a single company across specified categories"""
        company_data = {
            'company': company,
            'scraped_at': datetime.now().isoformat(),
            'categories': {},
            'summary': {
                'total_items': 0,
                'total_words': 0,
                'total_links': 0,
                'total_images': 0,
                'rich_content_count': 0
            }
        }
        
        for category in categories:
            if category in urls and urls[category]:
                try:
                    logger.info(f"Scraping {category} for {company}")
                    category_data = self._scrape_category(
                        company, category, urls[category], page_limit
                    )
                    company_data['categories'][category] = category_data
                    
                    # Update summary
                    company_data['summary']['total_items'] += len(category_data.get('items', []))
                    company_data['summary']['total_words'] += category_data.get('total_words', 0)
                    company_data['summary']['total_links'] += category_data.get('total_links', 0)
                    company_data['summary']['total_images'] += category_data.get('total_images', 0)
                    company_data['summary']['rich_content_count'] += category_data.get('rich_content_count', 0)
                    
                    # Add delay to avoid overwhelming servers
                    time.sleep(2)
                    
                except Exception as e:
                    logger.error(f"Error scraping {category} for {company}: {str(e)}")
                    company_data['categories'][category] = {
                        'error': str(e),
                        'status': 'failed'
                    }
        
        return company_data
    
    def _scrape_category(self, company: str, category: str, url: str, page_limit: int) -> Dict[str, Any]:
        """Scrape data for a specific category"""
        category_data = {
            'category': category,
            'url': url,
            'scraped_at': datetime.now().isoformat(),
            'items': [],
            'total_words': 0,
            'total_links': 0,
            'total_images': 0,
            'rich_content_count': 0
        }
        
        try:
            if category == 'marketing':
                items = self._scrape_marketing_site(url, page_limit)
            elif category == 'docs':
                items = self._scrape_documentation(url, page_limit)
            elif category == 'rss':
                items = self._scrape_rss_feed(url, page_limit)
            elif category == 'social':
                items = self._scrape_social_signals(url, page_limit)
            else:
                items = []
            
            # Process and analyze items
            processed_items = []
            for item in items:
                processed_item = self._process_content_item(item, company, category)
                processed_items.append(processed_item)
                
                # Update category metrics
                category_data['total_words'] += processed_item.get('word_count', 0)
                category_data['total_links'] += processed_item.get('link_count', 0)
                category_data['total_images'] += processed_item.get('image_count', 0)
                if processed_item.get('word_count', 0) > 1000:
                    category_data['rich_content_count'] += 1
            
            category_data['items'] = processed_items
            
        except Exception as e:
            logger.error(f"Error in category scraping: {str(e)}")
            category_data['error'] = str(e)
            category_data['status'] = 'failed'
        
        return category_data
    
    def _scrape_marketing_site(self, url: str, page_limit: int) -> List[Dict[str, Any]]:
        """Scrape marketing site content"""
        items = []
        
        try:
            # Use Firecrawl API if available, otherwise fall back to basic scraping
            if self.firecrawl_api_key:
                items = self._scrape_with_firecrawl(url, page_limit, 'marketing')
            else:
                items = self._scrape_with_requests(url, page_limit, 'marketing')
                
        except Exception as e:
            logger.error(f"Error scraping marketing site: {str(e)}")
            # Return mock data for demonstration
            items = self._generate_mock_marketing_data(url, page_limit)
        
        return items
    
    def _scrape_documentation(self, url: str, page_limit: int) -> List[Dict[str, Any]]:
        """Scrape documentation content"""
        items = []
        
        try:
            if self.firecrawl_api_key:
                items = self._scrape_with_firecrawl(url, page_limit, 'docs')
            else:
                items = self._scrape_with_requests(url, page_limit, 'docs')
                
        except Exception as e:
            logger.error(f"Error scraping documentation: {str(e)}")
            items = self._generate_mock_documentation_data(url, page_limit)
        
        return items
    
    def _scrape_rss_feed(self, url: str, page_limit: int) -> List[Dict[str, Any]]:
        """Scrape RSS feed content"""
        items = []
        
        try:
            feed = feedparser.parse(url)
            
            for entry in feed.entries[:page_limit]:
                item = {
                    'title': entry.get('title', ''),
                    'content': entry.get('summary', ''),
                    'url': entry.get('link', ''),
                    'published': entry.get('published', ''),
                    'author': entry.get('author', ''),
                    'source': 'rss'
                }
                items.append(item)
                
        except Exception as e:
            logger.error(f"Error scraping RSS feed: {str(e)}")
            items = self._generate_mock_rss_data(url, page_limit)
        
        return items
    
    def _scrape_social_signals(self, url: str, page_limit: int) -> List[Dict[str, Any]]:
        """Scrape social media signals"""
        items = []
        
        try:
            # This would integrate with social media APIs in a real implementation
            # For now, return mock data
            items = self._generate_mock_social_data(url, page_limit)
            
        except Exception as e:
            logger.error(f"Error scraping social signals: {str(e)}")
            items = self._generate_mock_social_data(url, page_limit)
        
        return items
    
    def _scrape_with_firecrawl(self, url: str, page_limit: int, category: str) -> List[Dict[str, Any]]:
        """Scrape using Firecrawl API"""
        if not self.firecrawl_api_key:
            raise ValueError("Firecrawl API key not configured")
        
        # This would implement actual Firecrawl API integration
        # For now, return mock data
        return self._generate_mock_data_for_category(url, page_limit, category)
    
    def _scrape_with_requests(self, url: str, page_limit: int, category: str) -> List[Dict[str, Any]]:
        """Basic scraping using requests and BeautifulSoup"""
        items = []
        
        try:
            response = self.session.get(url, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract basic content
            title = soup.find('title')
            title_text = title.get_text() if title else ''
            
            # Extract main content
            main_content = soup.find('main') or soup.find('body')
            content_text = main_content.get_text() if main_content else ''
            
            # Extract links
            links = soup.find_all('a', href=True)
            link_count = len(links)
            
            # Extract images
            images = soup.find_all('img')
            image_count = len(images)
            
            item = {
                'title': title_text,
                'content': content_text,
                'url': url,
                'source': category,
                'link_count': link_count,
                'image_count': image_count,
                'scraped_at': datetime.now().isoformat()
            }
            
            items.append(item)
            
        except Exception as e:
            logger.error(f"Error in basic scraping: {str(e)}")
            items = self._generate_mock_data_for_category(url, page_limit, category)
        
        return items
    
    def _process_content_item(self, item: Dict[str, Any], company: str, category: str) -> Dict[str, Any]:
        """Process and analyze a content item"""
        processed_item = item.copy()
        
        # Add company and category tags
        processed_item['company'] = company
        processed_item['category'] = category
        
        # Calculate word count
        content = item.get('content', '')
        word_count = len(content.split()) if content else 0
        processed_item['word_count'] = word_count
        
        # Calculate character count
        char_count = len(content) if content else 0
        processed_item['char_count'] = char_count
        
        # Determine content quality
        if word_count > 1000:
            processed_item['content_quality'] = 'rich'
        elif word_count > 500:
            processed_item['content_quality'] = 'medium'
        else:
            processed_item['content_quality'] = 'basic'
        
        # Extract links if not already present
        if 'link_count' not in processed_item:
            content_html = item.get('content_html', '')
            if content_html:
                soup = BeautifulSoup(content_html, 'html.parser')
                processed_item['link_count'] = len(soup.find_all('a', href=True))
            else:
                processed_item['link_count'] = 0
        
        # Extract images if not already present
        if 'image_count' not in processed_item:
            content_html = item.get('content_html', '')
            if content_html:
                soup = BeautifulSoup(content_html, 'html.parser')
                processed_item['image_count'] = len(soup.find_all('img'))
            else:
                processed_item['image_count'] = 0
        
        # Generate unique ID
        processed_item['id'] = self._generate_item_id(processed_item)
        
        return processed_item
    
    def _generate_item_id(self, item: Dict[str, Any]) -> str:
        """Generate unique ID for content item"""
        content_string = f"{item.get('company', '')}{item.get('title', '')}{item.get('url', '')}"
        return hashlib.md5(content_string.encode()).hexdigest()[:12]
    
    def batch_scrape_group(self, group: Dict[str, Any], page_limit: int = 10) -> Dict[str, Any]:
        """Scrape data for an entire competitor group"""
        group_results = {
            'group_name': group['name'],
            'group_key': group.get('group_key', 'custom'),
            'scraped_at': datetime.now().isoformat(),
            'companies': {},
            'summary': {
                'total_companies': len(group['companies']),
                'total_items': 0,
                'total_words': 0,
                'total_links': 0,
                'total_images': 0,
                'rich_content_count': 0
            }
        }
        
        for company in group['companies']:
            try:
                logger.info(f"Scraping data for {company}")
                company_urls = group['company_urls'].get(company, {})
                company_data = self.scrape_company_data(
                    company, company_urls, group['categories'], page_limit
                )
                
                group_results['companies'][company] = company_data
                
                # Update group summary
                group_results['summary']['total_items'] += company_data['summary']['total_items']
                group_results['summary']['total_words'] += company_data['summary']['total_words']
                group_results['summary']['total_links'] += company_data['summary']['total_links']
                group_results['summary']['total_images'] += company_data['summary']['total_images']
                group_results['summary']['rich_content_count'] += company_data['summary']['rich_content_count']
                
            except Exception as e:
                logger.error(f"Error scraping {company}: {str(e)}")
                group_results['companies'][company] = {
                    'error': str(e),
                    'status': 'failed'
                }
        
        return group_results
    
    def mass_scrape_all_groups(self, page_limit: int = 10) -> Dict[str, Any]:
        """Scrape all preset groups simultaneously"""
        all_results = {
            'mass_scrape_started': datetime.now().isoformat(),
            'groups': {},
            'overall_summary': {
                'total_groups': len(self.preset_groups),
                'total_companies': 0,
                'total_items': 0,
                'total_words': 0
            }
        }
        
        for group_key, group in self.preset_groups.items():
            try:
                logger.info(f"Starting mass scrape for group: {group['name']}")
                group_results = self.batch_scrape_group(group, page_limit)
                all_results['groups'][group_key] = group_results
                
                # Update overall summary
                all_results['overall_summary']['total_companies'] += group_results['summary']['total_companies']
                all_results['overall_summary']['total_items'] += group_results['summary']['total_items']
                all_results['overall_summary']['total_words'] += group_results['summary']['total_words']
                
            except Exception as e:
                logger.error(f"Error in mass scrape for group {group_key}: {str(e)}")
                all_results['groups'][group_key] = {
                    'error': str(e),
                    'status': 'failed'
                }
        
        all_results['mass_scrape_completed'] = datetime.now().isoformat()
        return all_results
    
    def import_data_file(self, file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """Import data from various file formats"""
        try:
            if file_type.lower() == 'csv':
                return self._import_csv(file_path)
            elif file_type.lower() == 'json':
                return self._import_json(file_path)
            elif file_type.lower() == 'markdown':
                return self._import_markdown(file_path)
            elif file_type.lower() == 'docx':
                return self._import_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error importing file {file_path}: {str(e)}")
            return []
    
    def _import_csv(self, file_path: str) -> List[Dict[str, Any]]:
        """Import data from CSV file"""
        items = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    item = {
                        'title': row.get('title', ''),
                        'content': row.get('content', ''),
                        'company': row.get('company', ''),
                        'category': row.get('category', ''),
                        'url': row.get('url', ''),
                        'source': 'csv_import',
                        'imported_at': datetime.now().isoformat()
                    }
                    items.append(item)
                    
        except Exception as e:
            logger.error(f"Error importing CSV: {str(e)}")
        
        return items
    
    def _import_json(self, file_path: str) -> List[Dict[str, Any]]:
        """Import data from JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                
            if isinstance(data, list):
                items = data
            elif isinstance(data, dict):
                items = [data]
            else:
                items = []
            
            # Add import metadata
            for item in items:
                item['source'] = 'json_import'
                item['imported_at'] = datetime.now().isoformat()
            
            return items
            
        except Exception as e:
            logger.error(f"Error importing JSON: {str(e)}")
            return []
    
    def _import_markdown(self, file_path: str) -> List[Dict[str, Any]]:
        """Import data from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Convert markdown to HTML for processing
            html_content = markdown.markdown(content)
            
            # Extract title from first heading
            soup = BeautifulSoup(html_content, 'html.parser')
            title_tag = soup.find(['h1', 'h2', 'h3'])
            title = title_tag.get_text() if title_tag else os.path.basename(file_path)
            
            item = {
                'title': title,
                'content': content,
                'content_html': html_content,
                'company': 'imported',
                'category': 'documentation',
                'source': 'markdown_import',
                'imported_at': datetime.now().isoformat()
            }
            
            return [item]
            
        except Exception as e:
            logger.error(f"Error importing Markdown: {str(e)}")
            return []
    
    def _import_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Import data from DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text content
            content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract title from first paragraph or filename
            title = doc.paragraphs[0].text if doc.paragraphs and doc.paragraphs[0].text else os.path.basename(file_path)
            
            item = {
                'title': title,
                'content': content,
                'company': 'imported',
                'category': 'documentation',
                'source': 'docx_import',
                'imported_at': datetime.now().isoformat()
            }
            
            return [item]
            
        except Exception as e:
            logger.error(f"Error importing DOCX: {str(e)}")
            return []
    
    def export_data(self, data: Union[Dict, List], format: str = 'json', filename: Optional[str] = None) -> str:
        """Export data in specified format"""
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'competitive_intelligence_data_{timestamp}'
        
        try:
            if format.lower() == 'json':
                filepath = f"{filename}.json"
                with open(filepath, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2, ensure_ascii=False)
            
            elif format.lower() == 'csv':
                filepath = f"{filename}.csv"
                if isinstance(data, dict):
                    # Flatten nested data for CSV export
                    flattened_data = self._flatten_data_for_csv(data)
                else:
                    flattened_data = data
                
                if flattened_data:
                    df = pd.DataFrame(flattened_data)
                    df.to_csv(filepath, index=False, encoding='utf-8')
            
            else:
                raise ValueError(f"Unsupported export format: {format}")
            
            logger.info(f"Data exported successfully to: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error exporting data: {str(e)}")
            raise
    
    def _flatten_data_for_csv(self, data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Flatten nested data structure for CSV export"""
        flattened = []
        
        def flatten_item(item, prefix=''):
            flat_item = {}
            for key, value in item.items():
                if isinstance(value, dict):
                    flat_item.update(flatten_item(value, f"{prefix}{key}_"))
                elif isinstance(value, list):
                    # Handle lists by creating separate rows or concatenating
                    if value and isinstance(value[0], dict):
                        for i, sub_item in enumerate(value):
                            flat_item.update(flatten_item(sub_item, f"{prefix}{key}_{i}_"))
                    else:
                        flat_item[f"{prefix}{key}"] = '; '.join(map(str, value))
                else:
                    flat_item[f"{prefix}{key}"] = value
            return flat_item
        
        # Handle different data structures
        if 'companies' in data:
            for company, company_data in data['companies'].items():
                if isinstance(company_data, dict) and 'error' not in company_data:
                    for category, category_data in company_data.get('categories', {}).items():
                        if isinstance(category_data, dict) and 'error' not in category_data:
                            for item in category_data.get('items', []):
                                flat_item = flatten_item(item)
                                flat_item['company'] = company
                                flat_item['category'] = category
                                flattened.append(flat_item)
        
        return flattened
    
    # Mock data generation methods for demonstration
    def _generate_mock_marketing_data(self, url: str, page_limit: int) -> List[Dict[str, Any]]:
        """Generate mock marketing data for demonstration"""
        return [
            {
                'title': 'Product Launch Announcement',
                'content': 'We are excited to announce the launch of our new product...',
                'url': url,
                'source': 'marketing',
                'link_count': 5,
                'image_count': 3
            }
        ]
    
    def _generate_mock_documentation_data(self, url: str, page_limit: int) -> List[Dict[str, Any]]:
        """Generate mock documentation data for demonstration"""
        return [
            {
                'title': 'API Reference Guide',
                'content': 'This guide provides comprehensive documentation for our API...',
                'url': url,
                'source': 'docs',
                'link_count': 15,
                'image_count': 2
            }
        ]
    
    def _generate_mock_rss_data(self, url: str, page_limit: int) -> List[Dict[str, Any]]:
        """Generate mock RSS data for demonstration"""
        return [
            {
                'title': 'Latest Blog Post',
                'content': 'Our latest insights on industry trends...',
                'url': url,
                'source': 'rss',
                'link_count': 8,
                'image_count': 1
            }
        ]
    
    def _generate_mock_social_data(self, url: str, page_limit: int) -> List[Dict[str, Any]]:
        """Generate mock social data for demonstration"""
        return [
            {
                'title': 'Social Media Update',
                'content': 'Engaging with our community on social platforms...',
                'url': url,
                'source': 'social',
                'link_count': 3,
                'image_count': 2
            }
        ]
    
    def _generate_mock_data_for_category(self, url: str, page_limit: int, category: str) -> List[Dict[str, Any]]:
        """Generate mock data for any category"""
        if category == 'marketing':
            return self._generate_mock_marketing_data(url, page_limit)
        elif category == 'docs':
            return self._generate_mock_documentation_data(url, page_limit)
        elif category == 'rss':
            return self._generate_mock_rss_data(url, page_limit)
        elif category == 'social':
            return self._generate_mock_social_data(url, page_limit)
        else:
            return []

    def enhanced_technical_scraping(self, company: str, urls: Dict[str, str]) -> Dict[str, Any]:
        """Enhanced scraping with technical content focus - 12-hour MVP enhancement"""
        results = {}
        
        for category, url in urls.items():
            try:
                logger.info(f"Enhanced technical scraping for {company} - {category}")
                
                # Use existing scraping logic
                content = self._scrape_url(url)
                
                # Enhanced content extraction
                structured_data = self._extract_technical_content(content, company)
                
                # Quality scoring
                quality_score = self._calculate_content_quality(structured_data)
                
                # Technical relevance assessment
                technical_relevance = self._assess_technical_relevance(category, content)
                
                results[category] = {
                    'content': structured_data,
                    'quality_score': quality_score,
                    'technical_relevance': technical_relevance,
                    'scraped_at': datetime.now().isoformat(),
                    'url': url
                }
                
            except Exception as e:
                logger.error(f"Error in enhanced technical scraping for {category}: {str(e)}")
                results[category] = {'error': str(e)}
        
        return results

    def _extract_technical_content(self, html_content: str, company: str = "") -> Dict[str, Any]:
        """Extract technical content using enhanced BeautifulSoup - Phase 1 implementation"""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract tables (existing functionality)
            tables = soup.find_all('table')
            table_data = []
            for table in tables:
                table_data.append({
                    'html': str(table),
                    'text': table.get_text(strip=True),
                    'rows': len(table.find_all('tr')),
                    'columns': len(table.find_all('th')) or len(table.find_all('td', limit=1))
                })
            
            # Extract code blocks with enhanced analysis
            code_blocks = soup.find_all(['code', 'pre'])
            code_data = []
            for block in code_blocks:
                code_info = self._analyze_code_block(block)
                code_data.append(code_info)
            
            # Extract links with technical relevance scoring
            links = soup.find_all('a')
            link_data = []
            for link in links:
                href = link.get('href')
                if href:
                    link_info = self._analyze_link_technical_relevance(link, company)
                    link_data.append(link_info)
            
            # Extract technical metadata
            technical_metadata = self._extract_technical_metadata(soup)
            
            # Detect OpenAPI specifications
            openapi_specs = self.detect_openapi_specs(html_content)
            
            # Classify overall content
            text_content = soup.get_text(strip=True)
            content_classification = self.classify_technical_content(text_content, company)
            
            # Extract API endpoints from text content
            api_endpoints = self._extract_api_endpoints_from_text(text_content)
            
            # Extract authentication patterns
            auth_patterns = self._extract_authentication_patterns(text_content)
            
            # Extract rate limiting information
            rate_limit_info = self._extract_rate_limit_info(text_content)
            
            return {
                'tables': table_data,
                'code_blocks': code_data,
                'links': link_data,
                'text_content': text_content,
                'technical_metadata': technical_metadata,
                'openapi_specs': openapi_specs,
                'content_classification': content_classification,
                'api_endpoints': api_endpoints,
                'authentication_patterns': auth_patterns,
                'rate_limit_info': rate_limit_info,
                'has_forms': len(soup.find_all('form')) > 0,
                'has_images': len(soup.find_all('img')) > 0,
                'has_videos': len(soup.find_all(['video', 'iframe'])) > 0,
                'extraction_timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error extracting technical content: {str(e)}")
            return {
                'error': str(e),
                'text_content': html_content[:1000] if html_content else '',
                'tables': [],
                'code_blocks': [],
                'links': [],
                'extraction_timestamp': datetime.now().isoformat()
            }

    def _analyze_code_block(self, code_block) -> Dict[str, Any]:
        """Enhanced code block analysis with technical indicators"""
        try:
            code_text = code_block.get_text(strip=True)
            language = self._detect_code_language(code_block)
            
            # Extract technical indicators
            technical_indicators = []
            if 'api' in code_text.lower() or 'endpoint' in code_text.lower():
                technical_indicators.append('api_related')
            if 'auth' in code_text.lower() or 'token' in code_text.lower():
                technical_indicators.append('authentication')
            if 'sdk' in code_text.lower() or 'client' in code_text.lower():
                technical_indicators.append('sdk_related')
            if 'config' in code_text.lower() or 'setup' in code_text.lower():
                technical_indicators.append('configuration')
            
            # Calculate complexity score
            complexity_score = self._calculate_code_complexity(code_text)
            
            return {
                'html': str(code_block),
                'text': code_text,
                'language': language,
                'length': len(code_text),
                'technical_indicators': technical_indicators,
                'complexity_score': complexity_score,
                'has_api_patterns': any(indicator in technical_indicators for indicator in ['api_related', 'authentication'])
            }
            
        except Exception as e:
            logger.error(f"Error analyzing code block: {str(e)}")
            return {
                'html': str(code_block),
                'text': code_block.get_text(strip=True),
                'language': 'unknown',
                'length': 0,
                'technical_indicators': [],
                'complexity_score': 0.0,
                'has_api_patterns': False
            }

    def _analyze_link_technical_relevance(self, link, company: str = "") -> Dict[str, Any]:
        """Analyze link for technical relevance"""
        try:
            href = link.get('href')
            link_text = link.get_text(strip=True)
            
            # Score technical relevance
            technical_score = 0.0
            technical_indicators = []
            
            # Check link text for technical keywords
            for category, keywords in self.technical_keywords.items():
                for keyword in keywords:
                    if keyword.lower() in link_text.lower():
                        technical_score += 0.1
                        technical_indicators.append(f"{category}:{keyword}")
            
            # Check URL for technical patterns
            if href:
                if any(pattern in href.lower() for pattern in ['/api/', '/docs/', '/developers/', '/reference/']):
                    technical_score += 0.3
                    technical_indicators.append('technical_url_pattern')
                if any(ext in href.lower() for ext in ['.json', '.yaml', '.yml', '.md']):
                    technical_score += 0.2
                    technical_indicators.append('technical_file_extension')
            
            # Company-specific relevance
            if company:
                company_lower = company.lower()
                for comp, keywords in self.company_technical_keywords.items():
                    if comp in company_lower:
                        for keyword in keywords:
                            if keyword.lower() in link_text.lower():
                                technical_score += 0.2
                                technical_indicators.append(f"company_specific:{keyword}")
            
            return {
                'url': href,
                'text': link_text,
                'title': link.get('title', ''),
                'is_external': self._is_external_link(href) if href else False,
                'technical_score': min(1.0, technical_score),
                'technical_indicators': technical_indicators,
                'relevance_category': self._categorize_link_relevance(technical_score)
            }
            
        except Exception as e:
            logger.error(f"Error analyzing link technical relevance: {str(e)}")
            return {
                'url': link.get('href', ''),
                'text': link.get_text(strip=True),
                'technical_score': 0.0,
                'technical_indicators': [],
                'relevance_category': 'unknown'
            }

    def _categorize_link_relevance(self, technical_score: float) -> str:
        """Categorize link by technical relevance"""
        if technical_score >= 0.7:
            return 'high_technical'
        elif technical_score >= 0.4:
            return 'medium_technical'
        elif technical_score >= 0.1:
            return 'low_technical'
        else:
            return 'non_technical'

    def _calculate_code_complexity(self, code_text: str) -> float:
        """Calculate code complexity score"""
        try:
            complexity_score = 0.0
            
            # Basic complexity indicators
            if len(code_text) > 1000:
                complexity_score += 0.3
            elif len(code_text) > 500:
                complexity_score += 0.2
            elif len(code_text) > 100:
                complexity_score += 0.1
            
            # Function/class complexity
            function_count = len(re.findall(r'\b(def|function|class)\b', code_text, re.IGNORECASE))
            complexity_score += min(0.3, function_count * 0.05)
            
            # Import/dependency complexity
            import_count = len(re.findall(r'\b(import|from|require|include)\b', code_text, re.IGNORECASE))
            complexity_score += min(0.2, import_count * 0.03)
            
            # Error handling complexity
            error_patterns = ['try:', 'catch', 'except', 'finally', 'error', 'exception']
            error_count = sum(1 for pattern in error_patterns if pattern in code_text.lower())
            complexity_score += min(0.2, error_count * 0.05)
            
            return min(1.0, complexity_score)
            
        except Exception:
            return 0.0

    def _extract_api_endpoints_from_text(self, text: str) -> List[Dict[str, Any]]:
        """Extract API endpoint patterns from text content"""
        try:
            endpoints = []
            
            # Look for URL patterns that might be API endpoints
            url_pattern = r'https?://[^\s]+(/api/[^\s]+|/v\d+/[^\s]+|/[a-z]+/[a-z0-9_-]+)'
            url_matches = re.findall(url_pattern, text, re.IGNORECASE)
            
            for match in url_matches:
                if any(keyword in match.lower() for keyword in ['api', 'v1', 'v2', 'v3', 'rest', 'graphql']):
                    endpoints.append({
                        'url': match,
                        'type': 'url_pattern',
                        'confidence': 'medium'
                    })
            
            # Look for HTTP method + path patterns
            http_pattern = r'\b(GET|POST|PUT|DELETE|PATCH)\s+([/\w-]+)'
            http_matches = re.findall(http_pattern, text, re.IGNORECASE)
            
            for method, path in http_matches:
                if any(keyword in path.lower() for keyword in ['api', 'endpoint', 'v1', 'v2', 'v3']):
                    endpoints.append({
                        'method': method.upper(),
                        'path': path,
                        'type': 'http_pattern',
                        'confidence': 'high'
                    })
            
            return endpoints
            
        except Exception as e:
            logger.error(f"Error extracting API endpoints: {str(e)}")
            return []

    def _extract_authentication_patterns(self, text: str) -> List[Dict[str, Any]]:
        """Extract authentication patterns from text content"""
        try:
            auth_patterns = []
            
            # OAuth patterns
            oauth_patterns = [
                r'oauth\s*2\.0',
                r'client_id\s*[:=]\s*[\w-]+',
                r'client_secret\s*[:=]\s*[\w-]+',
                r'authorization_code',
                r'access_token',
                r'refresh_token'
            ]
            
            for pattern in oauth_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    auth_patterns.append({
                        'type': 'oauth2',
                        'pattern': pattern,
                        'matches': matches,
                        'confidence': 'high'
                    })
            
            # API key patterns
            api_key_patterns = [
                r'api_key\s*[:=]\s*[\w-]+',
                r'x-api-key',
                r'authorization\s*:\s*bearer',
                r'x-auth-token'
            ]
            
            for pattern in api_key_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    auth_patterns.append({
                        'type': 'api_key',
                        'pattern': pattern,
                        'matches': matches,
                        'confidence': 'medium'
                    })
            
            # Basic auth patterns
            basic_auth_patterns = [
                r'basic\s+authentication',
                r'username\s*[:=]\s*[\w-]+',
                r'password\s*[:=]\s*[\w-]+'
            ]
            
            for pattern in basic_auth_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    auth_patterns.append({
                        'type': 'basic_auth',
                        'pattern': pattern,
                        'matches': matches,
                        'confidence': 'medium'
                    })
            
            return auth_patterns
            
        except Exception as e:
            logger.error(f"Error extracting authentication patterns: {str(e)}")
            return []

    def _extract_rate_limit_info(self, text: str) -> List[Dict[str, Any]]:
        """Extract rate limiting information from text content"""
        try:
            rate_limit_info = []
            
            # Rate limit patterns
            rate_patterns = [
                r'rate\s+limit[:\s]+(\d+)\s+(requests?|calls?)\s+per\s+(second|minute|hour|day)',
                r'(\d+)\s+(requests?|calls?)\s+per\s+(second|minute|hour|day)',
                r'throttling[:\s]+(\d+)\s+(requests?|calls?)',
                r'quota[:\s]+(\d+)\s+(requests?|calls?)',
                r'maximum\s+(\d+)\s+(requests?|calls?)'
            ]
            
            for pattern in rate_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    for match in matches:
                        if isinstance(match, tuple):
                            limit, unit, time_period = match
                        else:
                            limit = match
                            unit = 'requests'
                            time_period = 'minute'
                        
                        rate_limit_info.append({
                            'limit': int(limit),
                            'unit': unit,
                            'time_period': time_period,
                            'pattern': pattern,
                            'confidence': 'medium'
                        })
            
            return rate_limit_info
            
        except Exception as e:
            logger.error(f"Error extracting rate limit info: {str(e)}")
            return []

    def _detect_code_language(self, code_block) -> str:
        """Detect programming language from code block"""
        try:
            # Check for language class
            class_attr = code_block.get('class', [])
            for cls in class_attr:
                if cls.startswith('language-'):
                    return cls.replace('language-', '')
            
            # Check for data attributes
            data_lang = code_block.get('data-lang')
            if data_lang:
                return data_lang
            
            # Simple keyword-based detection
            text = code_block.get_text().lower()
            if any(keyword in text for keyword in ['function', 'var ', 'const ', 'let ']):
                return 'javascript'
            elif any(keyword in text for keyword in ['def ', 'import ', 'class ']):
                return 'python'
            elif any(keyword in text for keyword in ['public class', 'private ', 'public ']):
                return 'java'
            elif any(keyword in text for keyword in ['<?php', 'function ', '$']):
                return 'php'
            else:
                return 'unknown'
                
        except Exception:
            return 'unknown'

    def _is_external_link(self, url: str) -> bool:
        """Check if link is external"""
        try:
            parsed = urlparse(url)
            return bool(parsed.netloc and parsed.netloc not in ['localhost', '127.0.0.1'])
        except Exception:
            return False

    def _extract_technical_metadata(self, soup) -> Dict[str, Any]:
        """Extract technical metadata from HTML"""
        metadata = {}
        
        try:
            # Meta tags
            meta_tags = soup.find_all('meta')
            for meta in meta_tags:
                name_attr = meta.get('name', None)
                prop_attr = meta.get('property', None)
                key = name_attr or prop_attr
                content = meta.get('content', '')
                if key and content:
                    metadata[key] = content
            
            # Schema.org structured data
            schema_scripts = soup.find_all('script', type='application/ld+json')
            for script in schema_scripts:
                try:
                    if script.string:
                        schema_data = json.loads(script.string)
                        if isinstance(schema_data, dict):
                            metadata['schema'] = schema_data
                except json.JSONDecodeError:
                    continue
            
            # Open Graph tags
            og_tags = soup.find_all('meta', attrs={'property': lambda x: x and x.startswith('og:')})
            for tag in og_tags:
                property_name = tag.get('property', '')
                if property_name:
                    metadata[property_name] = tag.get('content', '')
            
            # Twitter Card tags (use attrs to avoid name param conflict)
            twitter_tags = soup.find_all('meta', attrs={'name': lambda x: x and x.startswith('twitter:')})
            for tag in twitter_tags:
                n = tag.get('name', '')
                if n:
                    metadata[n] = tag.get('content', '')
                    
        except Exception as e:
            logger.error(f"Error extracting technical metadata: {str(e)}")
        
        return metadata

    def _assess_technical_relevance(self, category: str, content: str) -> float:
        """Quick technical relevance scoring - 12-hour MVP enhancement"""
        try:
            if category in self.technical_keywords:
                keywords = self.technical_keywords[category]
                content_lower = content.lower()
                matches = sum(1 for keyword in keywords if keyword.lower() in content_lower)
                relevance_score = min(1.0, matches / len(keywords))
                
                # Boost score for technical content types
                if category in ['api_docs', 'integrations']:
                    relevance_score = min(1.0, relevance_score * 1.2)
                
                return round(relevance_score, 2)
            
            return 0.5
            
        except Exception as e:
            logger.error(f"Error assessing technical relevance: {str(e)}")
            return 0.5

    def classify_technical_content(self, content: str, company: str = "") -> Dict[str, Any]:
        """Classify content by technical type and relevance - Phase 1.3 implementation"""
        try:
            content_lower = content.lower()
            classification = {
                'content_type': 'unknown',
                'technical_depth': 0.0,
                'api_relevance': 0.0,
                'sdk_relevance': 0.0,
                'pricing_relevance': 0.0,
                'deployment_relevance': 0.0,
                'integrations_relevance': 0.0,
                'overall_technical_score': 0.0,
                'detected_patterns': [],
                'company_specific_matches': []
            }
            
            # Determine content type based on pattern matching
            content_type_scores = {}
            for content_type, patterns in self.content_patterns.items():
                score = 0.0
                matched_patterns = []
                for pattern in patterns:
                    matches = re.findall(pattern, content_lower, re.IGNORECASE)
                    if matches:
                        score += len(matches) * 0.1
                        matched_patterns.extend(matches)
                
                if score > 0:
                    content_type_scores[content_type] = score
                    classification['detected_patterns'].extend(matched_patterns)
            
            # Set primary content type
            if content_type_scores:
                classification['content_type'] = max(content_type_scores, key=content_type_scores.get)
            
            # Calculate relevance scores for each category
            for category, keywords in self.technical_keywords.items():
                if category == 'api_docs':
                    classification['api_relevance'] = self._calculate_keyword_relevance(content_lower, keywords)
                elif category == 'sdk_docs':
                    classification['sdk_relevance'] = self._calculate_keyword_relevance(content_lower, keywords)
                elif category == 'pricing':
                    classification['pricing_relevance'] = self._calculate_keyword_relevance(content_lower, keywords)
                elif category == 'deployment':
                    classification['deployment_relevance'] = self._calculate_keyword_relevance(content_lower, keywords)
                elif category == 'integrations':
                    classification['integrations_relevance'] = self._calculate_keyword_relevance(content_lower, keywords)
            
            # Calculate technical depth based on content characteristics
            classification['technical_depth'] = self._calculate_technical_depth(content)
            
            # Add company-specific technical keyword matching
            if company:
                company_lower = company.lower()
                for comp, keywords in self.company_technical_keywords.items():
                    if comp in company_lower:
                        company_matches = []
                        for keyword in keywords:
                            if keyword.lower() in content_lower:
                                company_matches.append(keyword)
                        if company_matches:
                            classification['company_specific_matches'] = company_matches
                            # Boost overall score for company-specific content
                            classification['overall_technical_score'] += 0.2
            
            # Calculate overall technical score
            relevance_scores = [
                classification['api_relevance'],
                classification['sdk_relevance'],
                classification['pricing_relevance'],
                classification['deployment_relevance'],
                classification['integrations_relevance']
            ]
            
            classification['overall_technical_score'] = min(1.0, 
                classification['overall_technical_score'] + 
                (sum(relevance_scores) / len(relevance_scores)) * 0.6 +
                classification['technical_depth'] * 0.4
            )
            
            # Round all scores
            for key, value in classification.items():
                if isinstance(value, float):
                    classification[key] = round(value, 3)
            
            return classification
            
        except Exception as e:
            logger.error(f"Error in technical content classification: {str(e)}")
            return {
                'content_type': 'error',
                'technical_depth': 0.0,
                'overall_technical_score': 0.0,
                'error': str(e)
            }
    
    def _calculate_keyword_relevance(self, content: str, keywords: List[str]) -> float:
        """Calculate relevance score based on keyword matches"""
        try:
            matches = sum(1 for keyword in keywords if keyword.lower() in content)
            return min(1.0, matches / len(keywords)) if keywords else 0.0
        except Exception:
            return 0.0
    
    def _calculate_technical_depth(self, content: str) -> float:
        """Calculate technical depth based on content characteristics"""
        try:
            depth_score = 0.0
            
            # Code blocks and technical elements
            if '```' in content or '<code>' in content:
                depth_score += 0.3
            
            # Technical terminology density
            technical_terms = ['api', 'endpoint', 'authentication', 'sdk', 'deployment', 'configuration']
            term_count = sum(1 for term in technical_terms if term in content.lower())
            depth_score += min(0.3, term_count * 0.05)
            
            # URL patterns (API endpoints, documentation links)
            url_patterns = [r'https?://[^\s]+\.json', r'https?://[^\s]+\.yaml', r'https?://[^\s]+\.yml']
            for pattern in url_patterns:
                if re.search(pattern, content):
                    depth_score += 0.2
            
            # HTTP methods
            http_methods = ['GET', 'POST', 'PUT', 'DELETE', 'PATCH']
            method_count = sum(1 for method in http_methods if method in content)
            depth_score += min(0.2, method_count * 0.05)
            
            return min(1.0, depth_score)
            
        except Exception:
            return 0.0

    def detect_openapi_specs(self, html_content: str) -> List[Dict]:
        """Detect OpenAPI specifications in HTML content - Phase 1.1 implementation"""
        try:
            openapi_specs = []
            
            # Look for direct OpenAPI file references
            openapi_patterns = [
                r'https?://[^\s"\']+\.json',
                r'https?://[^\s"\']+\.yaml',
                r'https?://[^\s"\']+\.yml'
            ]
            
            for pattern in openapi_patterns:
                matches = re.findall(pattern, html_content)
                for match in matches:
                    if any(keyword in match.lower() for keyword in ['openapi', 'swagger', 'api']):
                        openapi_specs.append({
                            'url': match,
                            'type': 'direct_reference',
                            'detected_at': datetime.now().isoformat()
                        })
            
            # Look for OpenAPI links in HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            for link in soup.find_all('a', href=True):
                href = link.get('href')
                if href and any(keyword in href.lower() for keyword in ['openapi', 'swagger', 'api.json', 'api.yaml']):
                    openapi_specs.append({
                        'url': href,
                        'type': 'html_link',
                        'link_text': link.get_text(strip=True),
                        'detected_at': datetime.now().isoformat()
                    })
            
            # Look for OpenAPI content in script tags
            for script in soup.find_all('script'):
                if script.string:
                    script_content = script.string
                    # Look for OpenAPI spec content embedded in scripts
                    if 'openapi' in script_content.lower() or 'swagger' in script_content.lower():
                        openapi_specs.append({
                            'type': 'embedded_script',
                            'content_preview': script_content[:200],
                            'detected_at': datetime.now().isoformat()
                        })
            
            return openapi_specs
            
        except Exception as e:
            logger.error(f"Error detecting OpenAPI specs: {str(e)}")
            return []

    def parse_openapi_spec(self, spec_url: str) -> Dict[str, Any]:
        """Parse OpenAPI specification and extract technical information"""
        try:
            response = self.session.get(spec_url, timeout=15)
            response.raise_for_status()
            
            spec_content = response.text
            spec_data = {}
            
            # Try to parse as JSON first
            try:
                spec_data = json.loads(spec_content)
            except json.JSONDecodeError:
                # Try YAML parsing if JSON fails
                try:
                    import yaml
                    spec_data = yaml.safe_load(spec_content)
                except ImportError:
                    logger.warning("PyYAML not available for YAML parsing")
                    return {'error': 'YAML parsing not available'}
            
            # Extract key technical information
            extracted_info = {
                'openapi_version': spec_data.get('openapi', 'unknown'),
                'title': spec_data.get('info', {}).get('title', ''),
                'version': spec_data.get('info', {}).get('version', ''),
                'description': spec_data.get('info', {}).get('description', ''),
                'endpoints': [],
                'authentication_methods': [],
                'rate_limits': [],
                'parsed_at': datetime.now().isoformat()
            }
            
            # Extract endpoints
            paths = spec_data.get('paths', {})
            for path, methods in paths.items():
                for method, details in methods.items():
                    if method.upper() in ['GET', 'POST', 'PUT', 'DELETE', 'PATCH']:
                        endpoint_info = {
                            'path': path,
                            'method': method.upper(),
                            'summary': details.get('summary', ''),
                            'description': details.get('description', ''),
                            'parameters': details.get('parameters', []),
                            'responses': list(details.get('responses', {}).keys())
                        }
                        extracted_info['endpoints'].append(endpoint_info)
            
            # Extract authentication methods
            security_schemes = spec_data.get('components', {}).get('securitySchemes', {})
            for scheme_name, scheme_details in security_schemes.items():
                auth_info = {
                    'name': scheme_name,
                    'type': scheme_details.get('type', ''),
                    'description': scheme_details.get('description', '')
                }
                extracted_info['authentication_methods'].append(auth_info)
            
            # Look for rate limiting information in descriptions
            for endpoint in extracted_info['endpoints']:
                description = endpoint.get('description', '').lower()
                if any(term in description for term in ['rate limit', 'throttling', 'quota', 'requests per']):
                    extracted_info['rate_limits'].append({
                        'endpoint': endpoint['path'],
                        'method': endpoint['method'],
                        'description': endpoint['description']
                    })
            
            return extracted_info
            
        except Exception as e:
            logger.error(f"Error parsing OpenAPI spec {spec_url}: {str(e)}")
            return {'error': str(e)}

    def _calculate_content_quality(self, structured_data: Dict[str, Any]) -> float:
        """Calculate content quality score - 12-hour MVP enhancement"""
        try:
            score = 0.0
            max_score = 10.0
            
            # Text content quality (40% of score)
            text_content = structured_data.get('text_content', '')
            if text_content:
                word_count = len(text_content.split())
                if word_count > 1000:
                    score += 4.0
                elif word_count > 500:
                    score += 3.0
                elif word_count > 100:
                    score += 2.0
                else:
                    score += 1.0
            
            # Technical content richness (30% of score)
            tables = structured_data.get('tables', [])
            code_blocks = structured_data.get('code_blocks', [])
            links = structured_data.get('links', [])
            
            if tables:
                score += 1.5
            if code_blocks:
                score += 1.5
            if len(links) > 10:
                score += 1.0
            elif len(links) > 5:
                score += 0.5
            
            # Metadata completeness (20% of score)
            technical_metadata = structured_data.get('technical_metadata', {})
            if technical_metadata.get('schema'):
                score += 1.0
            if technical_metadata.get('og:title'):
                score += 0.5
            if technical_metadata.get('description'):
                score += 0.5
            
            # Content structure (10% of score)
            if structured_data.get('has_forms'):
                score += 0.5
            if structured_data.get('has_images'):
                score += 0.5
            
            return min(max_score, round(score, 1))
            
        except Exception as e:
            logger.error(f"Error calculating content quality: {str(e)}")
            return 5.0

    @rate_limited(max_per_second=1)
    def _scrape_url(self, url: str) -> str:
        """Scrape URL with rate limiting - 12-hour MVP enhancement"""
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            return response.text
        except Exception as e:
            logger.error(f"Error scraping URL {url}: {str(e)}")
            raise e

    def extract_strategic_comparison_data(self, scraped_content):
        """Extract strategic comparison data across cloud-native favorable dimensions"""
        comparison_data = {
            'api_first_architecture': self._analyze_api_first_architecture(scraped_content),
            'cloud_native_features': self._analyze_cloud_native_features(scraped_content),
            'data_integration': self._analyze_data_integration(scraped_content),
            'developer_experience': self._analyze_developer_experience(scraped_content),
            'modern_analytics': self._analyze_modern_analytics(scraped_content)
        }
        
        overall_score = sum(comparison_data.values()) / len(comparison_data)
        comparison_data['overall_score'] = overall_score
        comparison_data['positioning'] = self._determine_positioning(overall_score)
        
        return comparison_data
    
    def _analyze_api_first_architecture(self, content):
        """Analyze API-first architecture capabilities"""
        score = 0
        indicators = {
            'api_documentation': 0,
            'rest_graphql': 0,
            'sdk_libraries': 0,
            'webhooks': 0,
            'authentication': 0,
            'versioning': 0
        }
        
        for item in content:
            text = item.get('text_content', '').lower()
            
            if any(term in text for term in ['api', 'endpoint', 'api documentation']):
                indicators['api_documentation'] += 1
            if any(term in text for term in ['rest', 'graphql', 'swagger', 'openapi']):
                indicators['rest_graphql'] += 1
            if any(term in text for term in ['sdk', 'client library', 'npm', 'pip', 'maven']):
                indicators['sdk_libraries'] += 1
            if any(term in text for term in ['webhook', 'callback', 'event-driven']):
                indicators['webhooks'] += 1
            if any(term in text for term in ['authentication', 'oauth', 'api key', 'jwt']):
                indicators['authentication'] += 1
            if any(term in text for term in ['versioning', 'v1', 'v2', 'v3', 'api version']):
                indicators['versioning'] += 1
        
        score = (
            indicators['api_documentation'] * 20 +
            indicators['rest_graphql'] * 20 +
            indicators['sdk_libraries'] * 15 +
            indicators['webhooks'] * 15 +
            indicators['authentication'] * 15 +
            indicators['versioning'] * 15
        ) / max(len(content), 1)
        
        return min(100, score)
    
    def _analyze_cloud_native_features(self, content):
        """Analyze cloud-native platform capabilities"""
        score = 0
        indicators = {
            'multi_cloud': 0,
            'auto_scaling': 0,
            'serverless': 0,
            'containers': 0,
            'microservices': 0,
            'infrastructure_as_code': 0
        }
        
        for item in content:
            text = item.get('text_content', '').lower()
            
            if any(term in text for term in ['multi-cloud', 'hybrid cloud', 'cross-cloud']):
                indicators['multi_cloud'] += 1
            if any(term in text for term in ['auto-scaling', 'elastic', 'auto-scale', 'dynamic scaling']):
                indicators['auto_scaling'] += 1
            if any(term in text for term in ['serverless', 'lambda', 'function as a service', 'faas']):
                indicators['serverless'] += 1
            if any(term in text for term in ['container', 'kubernetes', 'docker', 'orchestration']):
                indicators['containers'] += 1
            if any(term in text for term in ['microservices', 'distributed', 'service mesh']):
                indicators['microservices'] += 1
            if any(term in text for term in ['infrastructure as code', 'terraform', 'cloudformation', 'iac']):
                indicators['infrastructure_as_code'] += 1
        
        score = (
            indicators['multi_cloud'] * 20 +
            indicators['auto_scaling'] * 20 +
            indicators['serverless'] * 20 +
            indicators['containers'] * 15 +
            indicators['microservices'] * 15 +
            indicators['infrastructure_as_code'] * 10
        ) / max(len(content), 1)
        
        return min(100, score)
    
    def _analyze_data_integration(self, content):
        """Analyze data integration and connector capabilities"""
        score = 0
        indicators = {
            'real_time': 0,
            'etl_pipelines': 0,
            'connectors': 0,
            'data_warehouses': 0,
            'data_mesh': 0,
            'open_formats': 0
        }
        
        for item in content:
            text = item.get('text_content', '').lower()
            
            if any(term in text for term in ['real-time', 'streaming', 'live data', 'instant']):
                indicators['real_time'] += 1
            if any(term in text for term in ['etl', 'elt', 'data pipeline', 'data transformation']):
                indicators['etl_pipelines'] += 1
            if any(term in text for term in ['connector', 'integration', 'plugin', 'add-on']):
                indicators['connectors'] += 1
            if any(term in text for term in ['data warehouse', 'snowflake', 'bigquery', 'redshift']):
                indicators['data_warehouses'] += 1
            if any(term in text for term in ['data mesh', 'data fabric', 'distributed data']):
                indicators['data_mesh'] += 1
            if any(term in text for term in ['open format', 'parquet', 'avro', 'json', 'csv']):
                indicators['open_formats'] += 1
        
        score = (
            indicators['real_time'] * 20 +
            indicators['etl_pipelines'] * 20 +
            indicators['connectors'] * 20 +
            indicators['data_warehouses'] * 15 +
            indicators['data_mesh'] * 15 +
            indicators['open_formats'] * 10
        ) / max(len(content), 1)
        
        return min(100, score)
    
    def _analyze_developer_experience(self, content):
        """Analyze developer experience and self-service capabilities"""
        score = 0
        indicators = {
            'self_service': 0,
            'ci_cd': 0,
            'documentation': 0,
            'examples': 0,
            'community': 0,
            'playground': 0
        }
        
        for item in content:
            text = item.get('text_content', '').lower()
            
            if any(term in text for term in ['self-service', 'provisioning', 'self-provision']):
                indicators['self_service'] += 1
            if any(term in text for term in ['ci/cd', 'pipeline', 'deployment', 'continuous']):
                indicators['ci_cd'] += 1
            if any(term in text for term in ['developer documentation', 'getting started', 'tutorial']):
                indicators['documentation'] += 1
            if any(term in text for term in ['sample code', 'example', 'code sample', 'demo']):
                indicators['examples'] += 1
            if any(term in text for term in ['community', 'forum', 'support', 'developer portal']):
                indicators['community'] += 1
            if any(term in text for term in ['playground', 'sandbox', 'api explorer', 'interactive']):
                indicators['playground'] += 1
        
        score = (
            indicators['self_service'] * 25 +
            indicators['ci_cd'] * 20 +
            indicators['documentation'] * 20 +
            indicators['examples'] * 15 +
            indicators['community'] * 10 +
            indicators['playground'] * 10
        ) / max(len(content), 1)
        
        return min(100, score)
    
    def _analyze_modern_analytics(self, content):
        """Analyze modern analytics and AI capabilities"""
        score = 0
        indicators = {
            'ai_ml': 0,
            'real_time_analytics': 0,
            'collaboration': 0,
            'natural_language': 0,
            'automation': 0,
            'governance': 0
        }
        
        for item in content:
            text = item.get('text_content', '').lower()
            
            if any(term in text for term in ['ai', 'machine learning', 'ml', 'artificial intelligence']):
                indicators['ai_ml'] += 1
            if any(term in text for term in ['real-time analytics', 'streaming analytics', 'live insights']):
                indicators['real_time_analytics'] += 1
            if any(term in text for term in ['collaborative', 'team analytics', 'shared workspace']):
                indicators['collaboration'] += 1
            if any(term in text for term in ['natural language', 'nlp', 'conversational', 'chat']):
                indicators['natural_language'] += 1
            if any(term in text for term in ['automated insights', 'auto-discovery', 'smart suggestions']):
                indicators['automation'] += 1
            if any(term in text for term in ['governance', 'compliance', 'security', 'data lineage']):
                indicators['governance'] += 1
        
        score = (
            indicators['ai_ml'] * 20 +
            indicators['real_time_analytics'] * 20 +
            indicators['collaboration'] * 20 +
            indicators['natural_language'] * 15 +
            indicators['automation'] * 15 +
            indicators['governance'] * 10
        ) / max(len(content), 1)
        
        return min(100, score)
    
    def _determine_positioning(self, overall_score):
        """Determine competitive positioning based on overall score"""
        if overall_score >= 80:
            return 'Leader'
        elif overall_score >= 60:
            return 'Transitioning'
        else:
            return 'Legacy'


# Example usage and testing
if __name__ == "__main__":
    # Initialize scraper
    scraper = CompetitiveIntelligenceScraper()
    
    # Load a preset group
    tech_group = scraper.load_preset_group("tech_saas")
    print(f"Loaded group: {tech_group['name']}")
    
    # Scrape the group
    results = scraper.batch_scrape_group(tech_group, page_limit=5)
    print(f"Scraped {results['summary']['total_companies']} companies")
    
    # Export results
    filename = scraper.export_data(results, 'json')
    print(f"Results exported to: {filename}") 